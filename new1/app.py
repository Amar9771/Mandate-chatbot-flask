import os
import datetime
from io import BytesIO
from flask import Flask, render_template, request, redirect, url_for, send_file
from sqlalchemy import create_engine
from sqlalchemy.orm import sessionmaker
from models import Base, Company, BDActivity, BWRClientData
import docx

# Explicitly define the templates folder path
template_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'templates')
app = Flask(__name__, template_folder=template_dir)

# Database setup
engine = create_engine('sqlite:///company_data.db', echo=True)
Base.metadata.create_all(engine)
Session = sessionmaker(bind=engine)

@app.route('/', methods=['GET', 'POST'])
def search_company():
    session = Session()
    company = None
    if request.method == 'POST':
        search_name = request.form['company_name']
        company = session.query(Company).filter(Company.CompanyName.ilike(f'%{search_name}%')).first()
    session.close()
    return render_template('search.html', company=company)

@app.route('/bd_activity/<company_name>', methods=['GET', 'POST'])
def bd_activity(company_name):
    session = Session()
    if request.method == 'POST':
        bd_name = request.form['bd_name']
        person_met = request.form['person_met']
        minutes = request.form['minutes']
        new_activity = BDActivity(
            CompanyName=company_name,
            BDName=bd_name,
            PersonMet=person_met,
            Minutes=minutes,
            EntryDate=datetime.datetime.now()
        )
        session.add(new_activity)
        session.commit()
        session.close()
        return redirect(url_for('search_company'))
    session.close()
    return render_template('bd_activity_form.html', company_name=company_name)

@app.route('/generate_profile/<company_name>')
def generate_profile(company_name):
    session = Session()
    company = session.query(Company).filter(Company.CompanyName == company_name).first()
    bd_activities = session.query(BDActivity).filter(BDActivity.CompanyName == company_name).all()
    bwr_data = session.query(BWRClientData).filter(BWRClientData.CompanyName == company_name).first()

    doc = docx.Document()
    doc.add_heading(f'Profile for {company.CompanyName}', 0)
    doc.add_paragraph(f'Address: {company.Address}')
    doc.add_paragraph(f'Email: {company.Email}')
    doc.add_paragraph(f'Phone: {company.Phone}')
    doc.add_paragraph(f'Brief: {company.Brief}')

    doc.add_heading('BD Activities', level=1)
    for bd in bd_activities:
        doc.add_paragraph(f'BD Name: {bd.BDName}, Person Met: {bd.PersonMet}, Minutes: {bd.Minutes}')

    doc.add_heading('BWR Client Data', level=1)
    if bwr_data:
        doc.add_paragraph(f'Escalation: {bwr_data.Escalation}')
        doc.add_paragraph(f'Alerts: {bwr_data.Alerts}')
        doc.add_paragraph(f'Issues: {bwr_data.Issues}')
        doc.add_paragraph(f'Rationale Link: {bwr_data.RationaleLink}')
        doc.add_paragraph(f'WD Requests: {bwr_data.WDRequests}')
        doc.add_paragraph(f'Rating History: {bwr_data.RatingHistory}')

    session.close()

    file_stream = BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)

    safe_filename = f"{company_name.replace(' ', '_')}_profile.docx"
    return send_file(file_stream,
                     download_name=safe_filename,
                     as_attachment=True)

if __name__ == '__main__':
    app.run(debug=True, port=8080)