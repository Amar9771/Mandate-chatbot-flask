from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Create a new Word document
doc = Document()

# Title
title = doc.add_heading('BWR Process Flowchart', level=1)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Description
doc.add_paragraph("This document illustrates the BWR process with a clean, professional two-lane flowchart layout.")

# Add two-column layout
table = doc.add_table(rows=1, cols=2)
table.autofit = True
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Operational Workflow'
hdr_cells[1].text = 'Approval & Publishing Workflow'

# Content for each workflow
operational_workflow = [
    "1. BD Brings Mandate",
    "2. BD Coordinators Register Mandate & Upload Supporting Documents",
    "3. DDP – Financial Capture",
    "4. Allocation",
    "5. Rating Team Validates the Data",
    "6. Update MIS System & Upload Any Supporting Documents",
    "7. Do Peer Review:\n   - Establish Sector Reviewers\n   - Capture BWR Projections\n   - Score Generation"
]

approval_workflow = [
    "1. Prepare Rating Note / Get Approval from SH",
    "2. QC (Quality Check) [AI Assistance Optional]",
    "3. Committee Review",
    "4. Prepare Minutes / Rationale Letter & Get Approval from SH",
    "5. QC (Second Level Check)",
    "6. Share Draft Rationale",
    "7. Publish"
]

# Add rows to the table
max_len = max(len(operational_workflow), len(approval_workflow))
for i in range(max_len):
    row_cells = table.add_row().cells
    row_cells[0].text = operational_workflow[i] if i < len(operational_workflow) else ''
    row_cells[1].text = approval_workflow[i] if i < len(approval_workflow) else ''

# Save the document
doc_path = "/mnt/data/BWR_Process_Flowchart.docx"
doc.save(doc_path)

doc_path
